import io

# Optional imports with robust fallbacks
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class FileExportGenerator:
    def __init__(self):
        pass

    def generate_docx(self, title: str, content: str) -> bytes:
        """
        Generates a Word Document (DOCX) representation of the content.
        Falls back to a rich-text text format if docx library is not loaded.
        """
        out_io = io.BytesIO()
        
        if DOCX_AVAILABLE:
            try:
                doc = docx.Document()
                doc.add_heading(title, level=1)
                
                # Split content into paragraphs and headers
                lines = content.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("# "):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith("## "):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith("### "):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith("- ") or line.startswith("* "):
                        doc.add_paragraph(line[2:], style='List Bullet')
                    else:
                        doc.add_paragraph(line)
                        
                doc.save(out_io)
                out_io.seek(0)
                return out_io.getvalue()
            except Exception as e:
                print(f"python-docx generation failed: {e}. Falling back to formatted raw txt.")
                
        # Fallback raw text export
        text_content = f"{title}\n" + "="*len(title) + f"\n\n{content}"
        return text_content.encode("utf-8")

    def generate_pdf(self, title: str, content: str) -> bytes:
        """
        Generates a PDF Document representation of the content.
        Falls back to a raw text file if reportlab is not loaded.
        """
        out_io = io.BytesIO()
        
        if PDF_AVAILABLE:
            try:
                doc = SimpleDocTemplate(out_io, pagesize=letter,
                                        rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
                styles = getSampleStyleSheet()
                
                # Modify sample Normal style for clean layout
                normal_style = styles["Normal"]
                normal_style.fontSize = 10
                normal_style.leading = 14
                normal_style.textColor = colors.HexColor("#2C2640")
                
                title_style = ParagraphStyle(
                    'ReportTitle',
                    parent=styles['Heading1'],
                    fontSize=22,
                    leading=26,
                    textColor=colors.HexColor("#4B3F9E"),
                    spaceAfter=20
                )
                
                h2_style = ParagraphStyle(
                    'ReportH2',
                    parent=styles['Heading2'],
                    fontSize=14,
                    leading=18,
                    textColor=colors.HexColor("#F0A742"),
                    spaceBefore=14,
                    spaceAfter=6
                )

                h3_style = ParagraphStyle(
                    'ReportH3',
                    parent=styles['Heading3'],
                    fontSize=11,
                    leading=15,
                    textColor=colors.HexColor("#1C1638"),
                    spaceBefore=10,
                    spaceAfter=4
                )

                story = []
                story.append(Paragraph(title, title_style))
                story.append(Spacer(1, 10))

                lines = content.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("# "):
                        story.append(Paragraph(line[2:], title_style))
                    elif line.startswith("## "):
                        story.append(Paragraph(line[3:], h2_style))
                    elif line.startswith("### "):
                        story.append(Paragraph(line[4:], h3_style))
                    elif line.startswith("- ") or line.startswith("* "):
                        bullet_style = ParagraphStyle(
                            'BulletItem',
                            parent=normal_style,
                            leftIndent=20,
                            firstLineIndent=-10,
                            spaceAfter=4
                        )
                        story.append(Paragraph(f"&bull; {line[2:]}", bullet_style))
                    else:
                        story.append(Paragraph(line, normal_style))
                        story.append(Spacer(1, 4))
                
                doc.build(story)
                out_io.seek(0)
                return out_io.getvalue()
            except Exception as e:
                print(f"reportlab PDF generation failed: {e}. Falling back to formatted raw txt.")
                
        # Fallback plain text format
        text_content = f"{title}\n" + "="*len(title) + f"\n\n{content}"
        return text_content.encode("utf-8")
